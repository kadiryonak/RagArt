"""Integration tests for /upload endpoint accepting multiple formats."""

from __future__ import annotations

import io
import json
from pathlib import Path

import pytest


@pytest.fixture
def client(tmp_path, monkeypatch):
    """Flask test client with a temp DATA_FOLDER so uploads don't pollute repo.

    Uploads land in the workspace's files/ dir, not the data root, so the
    test must inspect the workspace path.
    """
    import app as app_module
    monkeypatch.setattr(app_module.settings, "DATA_FOLDER", str(tmp_path), raising=False)
    # Re-init workspace manager pointed at tmp_path so each test is isolated,
    # and rebuild the RagRegistry against it (registry holds a manager ref).
    from src.workspaces import WorkspaceManager, DEFAULT_WORKSPACE_ID
    from src.services import RagRegistry
    app_module.workspace_manager = WorkspaceManager(str(tmp_path))
    app_module.rag_registry = RagRegistry(app_module.workspace_manager)
    app_module.app.config["TESTING"] = True

    # Tests assert "(tmp / 'name')"; expose the workspace files dir as that
    # location so they read the right place.
    ws_files = app_module.workspace_manager.files_dir(DEFAULT_WORKSPACE_ID)
    return app_module.app.test_client(), ws_files


def _post_file(c, filename: str, content: bytes):
    return c.post(
        "/upload",
        data={"file": (io.BytesIO(content), filename)},
        content_type="multipart/form-data",
    )


class TestJSONUpload:
    def test_valid_json_list(self, client):
        c, tmp = client
        payload = json.dumps([{"title": "Algoritma", "content": "Adım adım yöntem."}])
        r = _post_file(c, "algo.json", payload.encode("utf-8"))
        assert r.status_code == 200
        body = r.get_json()
        assert body["success"] is True
        assert body["format"] == "json"
        assert body["document_count"] == 1
        assert (tmp / "algo.json").exists()

    def test_invalid_json_rejected_and_cleaned_up(self, client):
        c, tmp = client
        r = _post_file(c, "broken.json", b"{ not json")
        assert r.status_code == 400
        # File must NOT be left in the data folder
        assert not (tmp / "broken.json").exists()


class TestTextUpload:
    def test_valid_txt(self, client):
        c, tmp = client
        content = "Bu bir Türkçe metin dosyasıdır.".encode("utf-8")
        r = _post_file(c, "note.txt", content)
        assert r.status_code == 200
        body = r.get_json()
        assert body["format"] == "text"  # TextLoader.name = "text"
        assert body["document_count"] == 1


class TestMarkdownUpload:
    def test_valid_md(self, client):
        c, tmp = client
        content = b"# Baslik\n\nBu yeterince uzun bir markdown icerigidir."
        r = _post_file(c, "doc.md", content)
        assert r.status_code == 200
        body = r.get_json()
        assert body["format"] == "markdown"
        assert (tmp / "doc.md").exists()


class TestPDFUpload:
    def test_valid_pdf(self, client):
        c, tmp = client
        # Build a tiny real PDF via reportlab if available
        reportlab = pytest.importorskip("reportlab")
        from reportlab.pdfgen import canvas
        buf = io.BytesIO()
        cv = canvas.Canvas(buf)
        cv.drawString(72, 720, "Algoritma adim adim yontem uygulamasidir.")
        cv.showPage()
        cv.save()

        r = _post_file(c, "small.pdf", buf.getvalue())
        assert r.status_code == 200
        body = r.get_json()
        assert body["format"] == "pdf"
        assert body["document_count"] == 1


class TestDOCXUpload:
    def test_valid_docx(self, client):
        c, tmp = client
        docx = pytest.importorskip("docx")
        from docx import Document as DocxDocument
        d = DocxDocument()
        d.add_paragraph("Algoritma adim adim uygulanan bir yontemdir.")
        d.add_paragraph("Python yorumlamali bir dildir.")
        buf = io.BytesIO()
        d.save(buf)

        r = _post_file(c, "doc.docx", buf.getvalue())
        assert r.status_code == 200
        body = r.get_json()
        assert body["format"] == "docx"
        assert body["document_count"] == 1


class TestUnsupportedFormat:
    def test_rejects_unknown_extension(self, client):
        c, tmp = client
        r = _post_file(c, "image.zip", b"PK\x03\x04 fake zip")
        assert r.status_code == 400
        assert "Allowed" in r.get_json()["error"] or "Unsupported" in r.get_json()["error"]


class TestListFiles:
    def test_lists_multiple_formats(self, client):
        c, tmp = client
        _post_file(c, "a.json", b'[{"x":"Algoritma adim adim yontemdir"}]')
        _post_file(c, "b.txt", b"Yeterince uzun bir test metin icerigi.")

        r = c.get("/list-files")
        assert r.status_code == 200
        body = r.get_json()
        assert body["total"] == 2
        assert "supported_extensions" in body
        # get_file_info returns the file extension stripped of '.' (json/txt/...)
        formats = {f["format"] for f in body["files"]}
        assert "json" in formats
        assert "txt" in formats  # get_file_info uses extension; list endpoint uses ext
