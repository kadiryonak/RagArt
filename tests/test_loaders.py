"""Unit tests for document loaders.

Strategy:
    - JSON / TXT / MD: fixture files written to tmp_path
    - PDF: runtime-generated via reportlab if available, else pypdf write
      (fallback minimal PDF bytes)
    - DOCX: runtime-generated via python-docx
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest

from src.loaders import (
    BaseLoader,
    DOCXLoader,
    JSONLoader,
    LoaderRegistry,
    MarkdownLoader,
    PDFLoader,
    TextLoader,
    get_default_registry,
)


# ----- JSONLoader -----


class TestJSONLoader:
    def test_list_format_produces_items(self, tmp_path):
        path = tmp_path / "algoritma.json"
        path.write_text(
            json.dumps([
                {"title": "Algoritma", "content": "Adım adım yöntem."},
                {"title": "Veri", "content": "Veri yapıları organizasyonu."},
            ], ensure_ascii=False),
            encoding="utf-8",
        )
        docs = JSONLoader().load(path)
        assert len(docs) == 2
        assert all(d.metadata["source"] == "algoritma.json" for d in docs)
        assert all(d.metadata["format"] == "json" for d in docs)
        assert docs[0].metadata["item_index"] == 0
        assert docs[1].metadata["item_index"] == 1
        assert "Adım adım" in docs[0].page_content

    def test_object_format_single_doc(self, tmp_path):
        path = tmp_path / "py.json"
        path.write_text(
            json.dumps({"title": "Python", "content": "Yorumlamalı dil."}),
            encoding="utf-8",
        )
        docs = JSONLoader().load(path)
        assert len(docs) == 1
        assert docs[0].metadata["item_index"] == 0
        assert "Python" in docs[0].page_content or "Yorumlamalı" in docs[0].page_content

    def test_min_length_skipped(self, tmp_path):
        path = tmp_path / "tiny.json"
        path.write_text('{"x":"a"}', encoding="utf-8")
        # "x: a" — 4 chars after extraction, below MIN_CONTENT_LEN
        assert JSONLoader().load(path) == []


# ----- TextLoader -----


class TestTextLoader:
    def test_utf8_text(self, tmp_path):
        path = tmp_path / "note.txt"
        path.write_text("Bu bir Türkçe metindir.\nİkinci satır.", encoding="utf-8")
        docs = TextLoader().load(path)
        assert len(docs) == 1
        assert "Türkçe" in docs[0].page_content
        assert docs[0].metadata["format"] == "txt"

    def test_latin1_fallback(self, tmp_path):
        path = tmp_path / "old.txt"
        # Write some bytes that aren't valid UTF-8
        path.write_bytes(b"caf\xe9 latin1 \xfctf-8 incompatible")
        docs = TextLoader().load(path)
        assert len(docs) == 1
        assert "latin1" in docs[0].page_content

    def test_too_short_skipped(self, tmp_path):
        path = tmp_path / "tiny.txt"
        path.write_text("hi", encoding="utf-8")
        assert TextLoader().load(path) == []


# ----- MarkdownLoader -----


class TestMarkdownLoader:
    def test_basic_markdown(self, tmp_path):
        path = tmp_path / "readme.md"
        path.write_text(
            "# Başlık\n\nBu bir paragraf.\n\n- liste maddesi 1\n- liste maddesi 2\n",
            encoding="utf-8",
        )
        docs = MarkdownLoader().load(path)
        assert len(docs) == 1
        assert "Başlık" in docs[0].page_content
        assert docs[0].metadata["format"] == "md"

    def test_code_fences_stripped(self, tmp_path):
        path = tmp_path / "code.md"
        path.write_text(
            "Başlık\n\n```python\ndef foo():\n    pass\n```\n\nDevam metni.",
            encoding="utf-8",
        )
        docs = MarkdownLoader().load(path)
        content = docs[0].page_content
        assert "def foo()" in content  # block içeriği korundu
        assert "```" not in content     # fence kaldırıldı

    def test_frontmatter_stripped(self, tmp_path):
        path = tmp_path / "post.md"
        path.write_text(
            "---\ntitle: X\ndate: 2026\n---\nGerçek içerik buraya geliyor "
            "ve yeterince uzun bir paragraf oluşturuyor.",
            encoding="utf-8",
        )
        docs = MarkdownLoader().load(path)
        assert len(docs) == 1
        assert "title: X" not in docs[0].page_content
        assert "Gerçek içerik" in docs[0].page_content


# ----- PDFLoader -----


def _make_pdf(path: Path, pages: list[str]):
    """Sayfa metinlerinden minimal PDF üret. reportlab varsa onu kullanır,
    yoksa testler skip edilir."""
    try:
        from reportlab.pdfgen import canvas
    except ImportError:
        pytest.skip("reportlab not installed; PDF integration test skipped")

    c = canvas.Canvas(str(path))
    for text in pages:
        # Tek bir TextObject hızlıca yeter
        c.drawString(72, 720, text)
        c.showPage()
    c.save()


class TestPDFLoader:
    def test_multi_page_pdf(self, tmp_path):
        path = tmp_path / "doc.pdf"
        _make_pdf(path, [
            "Algoritma birinci sayfa Turkce icerik.",
            "Python ikinci sayfa farkli icerik.",
            "Yapay zeka ucuncu sayfa.",
        ])
        docs = PDFLoader().load(path)
        assert len(docs) == 3
        assert all(d.metadata["format"] == "pdf" for d in docs)
        assert docs[0].metadata["page"] == 1
        assert docs[2].metadata["page"] == 3
        assert "Algoritma" in docs[0].page_content
        assert "Python" in docs[1].page_content

    def test_empty_page_skipped(self, tmp_path):
        # Boş (whitespace-only) sayfa atlanır
        path = tmp_path / "sparse.pdf"
        _make_pdf(path, [
            "Bu uzun bir Turkce metin icerigi paragrafidir.",
            " ",  # boş sayfa
            "Bu da uzun bir baska Turkce metin icerigi paragrafidir.",
        ])
        docs = PDFLoader().load(path)
        # Boş sayfa kayıt dışı — 3 sayfa ama 2 document
        assert len(docs) == 2

    def test_pypdf_missing_raises_clear_error(self, monkeypatch, tmp_path):
        # pypdf'in import edilemediği senaryoyu simüle et
        import builtins
        real_import = builtins.__import__

        def fake_import(name, *a, **kw):
            if name == "pypdf":
                raise ImportError("simulated absence")
            return real_import(name, *a, **kw)

        monkeypatch.setattr(builtins, "__import__", fake_import)
        path = tmp_path / "x.pdf"
        path.write_bytes(b"%PDF-1.4 fake")
        with pytest.raises(RuntimeError, match="pypdf is required"):
            PDFLoader().load(path)


# ----- DOCXLoader -----


def _make_docx(path: Path, paragraphs: list[str], table_rows: list[list[str]] | None = None):
    try:
        from docx import Document as DocxDocument
    except ImportError:
        pytest.skip("python-docx not installed; DOCX integration test skipped")

    doc = DocxDocument()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_rows:
        t = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for i, row in enumerate(table_rows):
            for j, cell in enumerate(row):
                t.cell(i, j).text = cell
    doc.save(str(path))


class TestDOCXLoader:
    def test_paragraphs_concatenated(self, tmp_path):
        path = tmp_path / "doc.docx"
        _make_docx(path, [
            "Algoritma adım adım uygulanan bir yöntemdir.",
            "Python yorumlamalı bir dildir.",
        ])
        docs = DOCXLoader().load(path)
        assert len(docs) == 1
        assert "Algoritma" in docs[0].page_content
        assert "Python" in docs[0].page_content
        assert docs[0].metadata["format"] == "docx"

    def test_table_flattened(self, tmp_path):
        path = tmp_path / "table.docx"
        _make_docx(path,
            ["Bu Turkce bir docx test dosyasidir paragraphlar var."],
            table_rows=[["Dil", "Yıl"], ["Python", "1991"], ["Java", "1995"]],
        )
        docs = DOCXLoader().load(path)
        content = docs[0].page_content
        assert "Python" in content
        assert "1991" in content
        assert "Java" in content
        assert "|" in content  # row separator


# ----- LoaderRegistry -----


class TestRegistry:
    def test_default_registry_has_all_loaders(self):
        r = get_default_registry()
        exts = r.supported_extensions
        assert ".json" in exts
        assert ".pdf" in exts
        assert ".docx" in exts
        assert ".md" in exts
        assert ".markdown" in exts
        assert ".txt" in exts

    def test_dispatches_by_extension(self, tmp_path):
        r = get_default_registry()
        p = tmp_path / "x.txt"
        p.write_text("Bu bir test metnidir.", encoding="utf-8")
        loader = r.get_loader(p)
        assert isinstance(loader, TextLoader)

    def test_unknown_extension(self, tmp_path):
        r = LoaderRegistry()
        p = tmp_path / "x.zip"
        assert r.get_loader(p) is None
        assert r.load(p) == []

    def test_register_custom_loader(self):
        r = LoaderRegistry()

        class FakeLoader(BaseLoader):
            name = "fake"
            extensions = {".fake"}
            def load(self, file_path):
                return []

        r.register(FakeLoader())
        assert ".fake" in r.supported_extensions
