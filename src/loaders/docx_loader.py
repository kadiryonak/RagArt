"""DOCXLoader — python-docx ile paragraph-bazlı extract.

Tek Document üretir (paragraph'lar birleştirilir). Çok büyük .docx'ler için
heading bazlı bölme yapılabilir ama v1'de tek belge yeterli — text splitter
zaten chunk'lar.

Tablolar: hücreler düz metin olarak satır satır eklenir (sıralı bilgi
korunur, ama yapı kaybolur). v2'de tablo-aware loader.
"""

from __future__ import annotations

from pathlib import Path
from typing import List

try:
    from langchain_core.documents import Document
except ImportError:
    from langchain.schema import Document

from src.loaders.base import BaseLoader


class DOCXLoader(BaseLoader):
    name = "docx"
    extensions = {".docx"}

    MIN_CONTENT = 20

    def load(self, file_path: Path) -> List[Document]:
        try:
            from docx import Document as DocxDocument  # python-docx
        except ImportError as e:
            raise RuntimeError(
                "python-docx is required for DOCX loading. "
                "Install with: pip install python-docx"
            ) from e

        doc = DocxDocument(str(file_path))

        parts: List[str] = []

        # Paragraflar
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if text:
                parts.append(text)

        # Tablolar — satır bazlı flatten
        for table in doc.tables:
            for row in table.rows:
                cells = [(c.text or "").strip() for c in row.cells]
                cells = [c for c in cells if c]
                if cells:
                    parts.append(" | ".join(cells))

        content = "\n".join(parts).strip()
        if len(content) < self.MIN_CONTENT:
            return []

        return [Document(
            page_content=content,
            metadata={
                "source": file_path.name,
                "format": "docx",
                "file_path": str(file_path),
                "item_index": 0,
            },
        )]
